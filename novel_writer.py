#!/usr/bin/env python3
"""AI Novel Writing Automation Tool."""

from __future__ import annotations

import argparse
import json
import os
import re
import subprocess
import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Callable

from anthropic import Anthropic
from dotenv import load_dotenv
from docx import Document
from openai import OpenAI
from rich.console import Console
from rich.panel import Panel

console = Console()

DEFAULT_CONFIG = {
    "provider": "openai",
    "model": "gpt-4o",
    "project_dir": "./my_novel_project",
    "book_title": "My Book Title",
    "genre": "Epic Fantasy",
}

DEFAULT_PROGRESS = {
    "step1_brainstorm": False,
    "step2_story_dossier": False,
    "step3_characters": False,
    "step4_worldbuilding": False,
    "step5_outline": False,
    "step6_style_sheet": False,
    "scenes": {},
}


STEP_SPECS = {
    1: {
        "name": "Brainstorm",
        "output": "outputs/step1_brainstorm.docx",
    },
    2: {
        "name": "Story Dossier",
        "output": "outputs/step2_story_dossier.docx",
        "inputs": ["inputs/braindump.docx", "inputs/genre_tropes.docx"],
    },
    3: {
        "name": "Characters",
        "output": "outputs/step3_characters.docx",
        "inputs": [
            "inputs/braindump.docx",
            "outputs/step2_story_dossier.docx",
            "inputs/genre_tropes.docx",
        ],
    },
    4: {
        "name": "Worldbuilding",
        "output": "outputs/step4_worldbuilding.docx",
        "inputs": [
            "inputs/braindump.docx",
            "outputs/step2_story_dossier.docx",
            "inputs/genre_tropes.docx",
        ],
    },
    5: {
        "name": "Outline",
        "output": "outputs/step5_outline.docx",
        "inputs": [
            "inputs/braindump.docx",
            "outputs/step3_characters.docx",
            "outputs/step4_worldbuilding.docx",
            "inputs/genre_tropes.docx",
            "inputs/outline_template.docx",
        ],
    },
    6: {
        "name": "Style Sheet",
        "output": "outputs/step6_style_sheet.docx",
        "inputs": ["inputs/writing_samples.docx"],
    },
}


@dataclass
class RuntimeConfig:
    provider: str
    model: str
    project_dir: Path
    book_title: str
    genre: str


def build_xml_tag(tag: str, content: str) -> str:
    return f"<{tag}>\n{content}\n</{tag}>"


def load_docx_text(filepath: str | Path) -> str:
    path = Path(filepath)
    if not path.exists():
        raise FileNotFoundError(
            f"Required file not found: {path}. Please create/place this file before running this step."
        )
    document = Document(path)
    return "\n".join(p.text for p in document.paragraphs).strip()


def save_as_docx(text: str, filepath: str | Path) -> None:
    path = Path(filepath)
    path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    for raw_line in text.splitlines():
        line = raw_line.rstrip()
        if line.startswith("### "):
            doc.add_paragraph(line[4:].strip(), style="Heading 2")
        elif line.startswith("## "):
            doc.add_paragraph(line[3:].strip(), style="Heading 1")
        else:
            doc.add_paragraph(line, style="Normal")
    doc.save(path)


def open_file_and_wait(filepath: str | Path) -> None:
    path = str(filepath)
    if sys.platform == "darwin":
        subprocess.run(["open", path], check=False)
    elif os.name == "nt":
        os.startfile(path)  # type: ignore[attr-defined]
    else:
        subprocess.run(["xdg-open", path], check=False)

    console.print(f"\n✏️  File opened: {path}\nEdit it now, then press ENTER when done...")
    input()


def _extract_anthropic_text(resp) -> str:
    parts = []
    for block in resp.content:
        if getattr(block, "type", None) == "text":
            parts.append(block.text)
    return "\n".join(parts).strip()


def call_ai(system_prompt: str, user_prompt: str, config: RuntimeConfig) -> str:
    provider = config.provider.lower()

    try:
        if provider in {"openai", "openrouter"}:
            if provider == "openai":
                api_key = os.getenv("OPENAI_API_KEY")
                base_url = None
            else:
                api_key = os.getenv("OPENROUTER_API_KEY")
                base_url = "https://openrouter.ai/api/v1"

            if not api_key:
                raise RuntimeError(
                    f"Missing API key for provider '{provider}'. Please set the correct key in .env."
                )

            client = OpenAI(api_key=api_key, base_url=base_url)
            messages = []
            if system_prompt.strip():
                messages.append({"role": "system", "content": system_prompt})
            messages.append({"role": "user", "content": user_prompt})
            response = client.chat.completions.create(
                model=config.model,
                messages=messages,
                max_tokens=8192,
            )
            return (response.choices[0].message.content or "").strip()

        if provider == "anthropic":
            api_key = os.getenv("ANTHROPIC_API_KEY")
            if not api_key:
                raise RuntimeError("Missing ANTHROPIC_API_KEY in environment/.env.")
            client = Anthropic(api_key=api_key)
            response = client.messages.create(
                model=config.model,
                max_tokens=8192,
                system=system_prompt or "",
                messages=[{"role": "user", "content": user_prompt}],
            )
            return _extract_anthropic_text(response)

        raise ValueError("Unsupported provider. Use openai, anthropic, or openrouter.")
    except Exception as exc:
        raise RuntimeError(f"AI call failed: {exc}") from exc


def slugify_chapter_name(name: str) -> str:
    slug = name.lower().strip()
    slug = re.sub(r"[^a-z0-9]+", "_", slug)
    slug = re.sub(r"_+", "_", slug).strip("_")
    return slug or "chapter"


def load_json(path: Path, default_data: dict) -> dict:
    if not path.exists():
        return default_data.copy()
    return json.loads(path.read_text(encoding="utf-8"))


def save_json(path: Path, data: dict) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(data, indent=2), encoding="utf-8")


def ensure_required_file(project_dir: Path, rel_path: str) -> Path:
    path = project_dir / rel_path
    if not path.exists():
        console.print(
            f"[yellow]Missing required file:[/yellow] {rel_path}\n"
            f"Please place it at: {path}"
        )
        raise SystemExit(1)
    return path


def ai_with_retry(system_prompt: str, user_prompt: str, config: RuntimeConfig) -> str:
    with console.status(
        f"⏳ Calling {config.provider} ({config.model})...", spinner="dots"
    ):
        try:
            return call_ai(system_prompt, user_prompt, config)
        except RuntimeError as first_err:
            console.print(f"[red]{first_err}[/red]")

    retry = input("Retry? (y/n): ").strip().lower() == "y"
    if not retry:
        raise SystemExit(1)

    with console.status(
        f"⏳ Retrying {config.provider} ({config.model})...", spinner="dots"
    ):
        try:
            return call_ai(system_prompt, user_prompt, config)
        except RuntimeError as second_err:
            console.print(f"[red]{second_err}[/red]")
            raise SystemExit(1)


def run_step_1(config: RuntimeConfig, project_dir: Path, skip_edit: bool) -> Path:
    system = (
        "You are an expert creative writer and outliner. Your task is to help a human take their ideas "
        "and build it into an excellent series outline. It is important, as you are brainstorming, that "
        "the ideas remain consistent with the tropes and conventions of the current genre.\n"
        f"The current genre is: {config.genre}"
    )
    user = (
        "I'd like to brainstorm ideas for my novel. Please help me develop a strong series concept, "
        "including premise, core conflict, key characters, and major plot beats. Ask me questions if you "
        "need more details, or generate 5-10 ideas for any element you think is missing."
    )

    output = project_dir / "outputs/step1_brainstorm.docx"
    result = ai_with_retry(system, user, config)
    save_as_docx(result, output)
    console.print(f"✅ Saved: {output.relative_to(project_dir)}")
    if not skip_edit:
        open_file_and_wait(output)
    return output


def run_step_2(config: RuntimeConfig, project_dir: Path, skip_edit: bool) -> Path:
    braindump = load_docx_text(ensure_required_file(project_dir, "inputs/braindump.docx"))
    tropes = load_docx_text(ensure_required_file(project_dir, "inputs/genre_tropes.docx"))

    user = "\n\n".join(
        [
            build_xml_tag("braindump", braindump),
            build_xml_tag("genre_tropes", tropes),
            build_xml_tag(
                "instructions",
                (
                    "Given the above braindump information and genre tropes, I want you to create a "
                    "pre-writing story dossier that has a list of everything the author will need to fully "
                    f"flesh out the characters, worldbuilding, and outline for {config.book_title}. "
                    "Here's what you should include:\n\n"
                    "Characters: Make a complete list of characters, including minor ones, needed for this "
                    "book. Label these as appropriate with their role in the story (protagonist, antagonist, "
                    "side character, henchman, comic relief, love interest, etc.) Give a brief explanation of "
                    "who the character is and their role in the story in no more than 1-2 sentences per "
                    "character. When referencing a group of individuals, name 3-5 minor characters who are "
                    "part of that group.\n\n"
                    "Worldbuilding Info: Make a complete list of all locations, objects, and other worldbuilding "
                    "information that will be needed for this book. Give a brief explanation of each element in "
                    "no more than one sentence.\n\n"
                    "Outline: Make a plan on what will be needed to fully outline this book. Do not start "
                    "outlining yet — just give instructions and suggestions on what will be needed."
                ),
            ),
        ]
    )

    output = project_dir / "outputs/step2_story_dossier.docx"
    result = ai_with_retry("", user, config)
    save_as_docx(result, output)
    console.print(f"✅ Saved: {output.relative_to(project_dir)}")
    if not skip_edit:
        open_file_and_wait(output)
    return output


def run_step_3(config: RuntimeConfig, project_dir: Path, skip_edit: bool) -> Path:
    braindump = load_docx_text(ensure_required_file(project_dir, "inputs/braindump.docx"))
    dossier = load_docx_text(ensure_required_file(project_dir, "outputs/step2_story_dossier.docx"))
    tropes = load_docx_text(ensure_required_file(project_dir, "inputs/genre_tropes.docx"))

    user = "\n\n".join(
        [
            build_xml_tag("braindump", braindump),
            build_xml_tag("dossier", dossier),
            build_xml_tag("genre_tropes", tropes),
            build_xml_tag(
                "instructions",
                (
                    "Given the above braindump information, genre tropes, and outline dossier, I want you to "
                    f"create a fleshed-out list of all the important characters needed specifically for {config.book_title}.\n\n"
                    "For each major character, include:\n"
                    "1. A physical description\n"
                    "2. Their primary role in the story\n"
                    "3. Their most appropriate Myers-Briggs profile, Enneagram, and Clifton Strengths\n"
                    "4. Their core motivation\n"
                    "5. A brief background before the start of the story\n"
                    "6. An interesting quirk, hobby, or unique trait\n"
                    "7. Dialogue style\n"
                    "8. Dialogue examples in relaxed, stressful, thoughtful, and exciting situations\n\n"
                    "Format using Markdown:\n"
                    f"## {config.book_title}\n"
                    "### Major Character Name:\n"
                    "* Physical Description:\n"
                    "* Role in Story:\n"
                    "* Personality Profiles:\n"
                    "* Core Motivation:\n"
                    "* Background:\n"
                    "* Quirk:\n"
                    "* Dialogue Style:\n"
                    "* Dialogue Samples:\n\n"
                    "### Minor Characters:\n"
                    "* [NAME]: [1-2 sentences]\n\n"
                    "Only include the asked-for character details. No preamble or commentary."
                ),
            ),
        ]
    )

    output = project_dir / "outputs/step3_characters.docx"
    result = ai_with_retry("", user, config)
    save_as_docx(result, output)
    console.print(f"✅ Saved: {output.relative_to(project_dir)}")
    if not skip_edit:
        open_file_and_wait(output)
    return output


def run_step_4(config: RuntimeConfig, project_dir: Path, skip_edit: bool) -> Path:
    braindump = load_docx_text(ensure_required_file(project_dir, "inputs/braindump.docx"))
    dossier = load_docx_text(ensure_required_file(project_dir, "outputs/step2_story_dossier.docx"))
    tropes = load_docx_text(ensure_required_file(project_dir, "inputs/genre_tropes.docx"))

    user = "\n\n".join(
        [
            build_xml_tag("braindump", braindump),
            build_xml_tag("dossier", dossier),
            build_xml_tag("genre_tropes", tropes),
            build_xml_tag(
                "instructions",
                (
                    "Given the above braindump, genre tropes, and dossier, create a fleshed-out list of all "
                    f"important setting and worldbuilding elements needed specifically for {config.book_title}.\n\n"
                    "Organize worldbuilding elements into applicable categories. Possible categories include:\n"
                    "- High-level Worldbuilding\n- Settings/Locations\n- Magic Systems/Technology\n- Groups/Races\n"
                    "- Gods/Deities\n- Geography/Nature\n- Population/Politics\n- Culture\n"
                    "- History/Lore\n- Religion/Beliefs\n- Languages\n\n"
                    "Only use categories that apply and have content. Format using Markdown:\n\n"
                    f"## {config.book_title}\n"
                    "### WORLDBUILDING CATEGORY:\n"
                    "* NAME OF ELEMENT: [3-4 specific sentences]\n\n"
                    "Only include the asked-for worldbuilding details. No preamble or commentary."
                ),
            ),
        ]
    )

    output = project_dir / "outputs/step4_worldbuilding.docx"
    result = ai_with_retry("", user, config)
    save_as_docx(result, output)
    console.print(f"✅ Saved: {output.relative_to(project_dir)}")
    if not skip_edit:
        open_file_and_wait(output)
    return output


def run_step_5(config: RuntimeConfig, project_dir: Path, skip_edit: bool) -> Path:
    worldbuilding = load_docx_text(ensure_required_file(project_dir, "outputs/step4_worldbuilding.docx"))
    characters = load_docx_text(ensure_required_file(project_dir, "outputs/step3_characters.docx"))
    braindump = load_docx_text(ensure_required_file(project_dir, "inputs/braindump.docx"))
    tropes = load_docx_text(ensure_required_file(project_dir, "inputs/genre_tropes.docx"))
    outline_template = load_docx_text(ensure_required_file(project_dir, "inputs/outline_template.docx"))

    user = "\n\n".join(
        [
            build_xml_tag("worldbuilding", worldbuilding),
            build_xml_tag("characters", characters),
            build_xml_tag("braindump", braindump),
            build_xml_tag("genre_tropes", tropes),
            build_xml_tag("outline_template", outline_template),
            build_xml_tag(
                "instructions",
                (
                    f"Using the above information, generate a fully fleshed-out outline for {config.book_title}. "
                    "Each chapter summary should be specific rather than vague — write as if handing it "
                    "to a ghostwriter.\n\n"
                    "Format using Markdown:\n"
                    f"## {config.book_title}\n"
                    "### CHAPTER TITLE:\n"
                    "[200-250 word description in 1-3 paragraphs with specific details]\n\n"
                    "Only include the asked-for outline details. No preamble or commentary."
                ),
            ),
        ]
    )

    output = project_dir / "outputs/step5_outline.docx"
    result = ai_with_retry("", user, config)
    save_as_docx(result, output)
    console.print(f"✅ Saved: {output.relative_to(project_dir)}")
    if not skip_edit:
        open_file_and_wait(output)
    return output


def run_step_6(config: RuntimeConfig, project_dir: Path, skip_edit: bool) -> Path:
    writing_samples = load_docx_text(ensure_required_file(project_dir, "inputs/writing_samples.docx"))
    user = "\n\n".join(
        [
            build_xml_tag("writing_samples", writing_samples),
            (
                "Given the above writing samples, draft a prose style sheet with instructions on how to write "
                "like these samples, including small snippets as examples. This style sheet is for use by an "
                "LLM to write fiction in the same style.\n\n"
                "Requirements:\n"
                "- Emphasize deep point of view and show-don't-tell, with examples from the samples\n"
                "- Specify POV and tenses used\n"
                "- Include a section on dialogue style\n"
                "- Include a section on average grade-level of the text\n"
                "- Do not focus on specific characters or plot — only on prose style\n"
                "- Be thorough"
            ),
        ]
    )

    output = project_dir / "outputs/step6_style_sheet.docx"
    result = ai_with_retry("", user, config)
    save_as_docx(result, output)
    console.print(f"✅ Saved: {output.relative_to(project_dir)}")
    if not skip_edit:
        open_file_and_wait(output)
    return output


def run_scene_brief(
    config: RuntimeConfig,
    project_dir: Path,
    chapter_name: str,
    previous_chapter_path: str | None,
    skip_edit: bool,
) -> Path:
    outline = load_docx_text(ensure_required_file(project_dir, "outputs/step5_outline.docx"))
    characters = load_docx_text(ensure_required_file(project_dir, "outputs/step3_characters.docx"))
    worldbuilding = load_docx_text(ensure_required_file(project_dir, "outputs/step4_worldbuilding.docx"))
    previous = ""
    if previous_chapter_path:
        previous = load_docx_text(Path(previous_chapter_path))

    user = "\n\n".join(
        [
            build_xml_tag("outline", outline),
            build_xml_tag("characters", characters),
            build_xml_tag("worldbuilding", worldbuilding),
            build_xml_tag("previous_chapter_text", previous),
            build_xml_tag(
                "instructions",
                (
                    f'Given the above outline and character/worldbuilding information, flesh out a "scene brief" for {chapter_name}. '
                    f'Label it clearly as "{chapter_name}". Include:\n\n'
                    "- POV: Third person limited — specify whose POV and justify the choice.\n"
                    f"- Genre: {config.genre}\n"
                    f'- Plot (Verbatim + Beats): Reproduce the full "{chapter_name}" plot summary verbatim from the outline, '
                    "then add 20-25 scene beats focused only on plot events (no sensory details).\n"
                    "- Scene Function: Define the narrative function (e.g., Inciting Incident, Character Introduction, etc.).\n"
                    "- Previous Chapter: Ensure plot picks up appropriately after the previous chapter. (Ignore if this is the first chapter.)\n"
                    "- Characters: List all characters in this chapter. For each: Name & Role, Physical Appearance (scene-specific), "
                    "Emotional State & Goals, Behavioural Notes.\n"
                    "- Setting: Sensory-rich description — time of day, terrain, sounds, smells, lighting, weather, tone.\n"
                    "- Main Source of Conflict: Central dramatic tension — internal, interpersonal, societal, or environmental.\n"
                    "- Tone & Style Notes: Specific voice cues, pacing guidance, dialogue intention.\n"
                    "- Symbolism or Thematic Layer: Symbols, metaphors, or archetypal moments to introduce.\n"
                    "- Continuity Considerations: Links to past chapters, foreshadowing, items and emotional threads that must remain consistent.\n"
                    "- Other Notes: Worldbuilding mechanics, scene transitions, or structural devices."
                ),
            ),
        ]
    )

    slug = slugify_chapter_name(chapter_name)
    output = project_dir / f"outputs/scenes/{slug}_brief.docx"
    result = ai_with_retry("", user, config)
    save_as_docx(result, output)
    console.print(f"✅ Saved: {output.relative_to(project_dir)}")
    if not skip_edit:
        open_file_and_wait(output)
    return output


def run_scene_draft(
    config: RuntimeConfig,
    project_dir: Path,
    chapter_name: str,
    chapter_slug: str,
    previous_chapter_path: str | None,
    skip_edit: bool,
) -> Path:
    samples = load_docx_text(ensure_required_file(project_dir, "inputs/writing_samples.docx"))
    style_sheet = load_docx_text(ensure_required_file(project_dir, "outputs/step6_style_sheet.docx"))
    prohibited = load_docx_text(ensure_required_file(project_dir, "inputs/prohibited_words.docx"))
    outline = load_docx_text(ensure_required_file(project_dir, "outputs/step5_outline.docx"))
    characters = load_docx_text(ensure_required_file(project_dir, "outputs/step3_characters.docx"))
    worldbuilding = load_docx_text(ensure_required_file(project_dir, "outputs/step4_worldbuilding.docx"))
    scene_brief = load_docx_text(ensure_required_file(project_dir, f"outputs/scenes/{chapter_slug}_brief.docx"))
    previous = ""
    if previous_chapter_path:
        previous = load_docx_text(Path(previous_chapter_path))

    full_context = "\n\n".join([outline, characters, worldbuilding])
    user = "\n\n".join(
        [
            build_xml_tag("prose_style_example", samples),
            build_xml_tag("style_sheet", style_sheet),
            build_xml_tag("prohibited_words", prohibited),
            build_xml_tag("full_context", full_context),
            build_xml_tag("previous_chapter_text", previous),
            build_xml_tag("scene_brief", scene_brief),
            build_xml_tag(
                "instructions",
                (
                    f'Write the entire "{chapter_name}" based on the scene brief. Cover it thoroughly from deep point of view, '
                    "written as if by a bestselling novelist. Use the prose_style_example to guide style.\n\n"
                    "Always follow these rules:\n"
                    "- Write in active voice\n"
                    "- Show, don't tell\n"
                    "- No adverbs, no clichés, no overused phrases\n"
                    "- Do not use any words from the prohibited_words list\n"
                    "- Convey events through dialogue; avoid \"he said/she said\" tags — show gestures/expressions instead\n"
                    "- Mix short punchy sentences with longer descriptive ones\n"
                    "- Dialogue gets its own paragraph\n"
                    "- Reduce uncertainty language (\"trying\", \"maybe\")\n"
                    "- Reduce metaphors\n"
                    "- Do not conclude the scene beyond what the scene brief specifies. Never end with foreshadowing beyond what is in the brief.\n"
                    "- Do not write further than instructed. Stop early if all beats are covered.\n"
                    "- Word count: approximately 3000 words\n"
                    "- If a previous chapter exists, begin where it left off with no overlap\n\n"
                    f"Format output as Markdown. Begin with the chapter name as an H2 heading: ## {chapter_name}"
                ),
            ),
        ]
    )

    output = project_dir / f"outputs/scenes/{chapter_slug}_draft.docx"
    result = ai_with_retry("", user, config)
    save_as_docx(result, output)
    console.print(f"✅ Saved: {output.relative_to(project_dir)}")
    if not skip_edit:
        open_file_and_wait(output)
    return output


def execute_step(step_num: int, config: RuntimeConfig, project_dir: Path, skip_edit: bool) -> Path:
    console.print(f"\n▶ Step {step_num}: {STEP_SPECS[step_num]['name']}")
    fn_map: dict[int, Callable[[RuntimeConfig, Path, bool], Path]] = {
        1: run_step_1,
        2: run_step_2,
        3: run_step_3,
        4: run_step_4,
        5: run_step_5,
        6: run_step_6,
    }
    return fn_map[step_num](config, project_dir, skip_edit)


def scene_flow(config: RuntimeConfig, project_dir: Path, chapter_name: str, skip_edit: bool) -> str:
    ask_prev = input("Is there a previous chapter draft to include? (y/n): ").strip().lower()
    previous_path = None
    if ask_prev == "y":
        previous_path = input("Enter previous chapter draft filepath: ").strip()

    console.print(f"\n▶ Step 7: Scene Brief — {chapter_name}")
    run_scene_brief(config, project_dir, chapter_name, previous_path, skip_edit)
    slug = slugify_chapter_name(chapter_name)

    console.print(f"\n▶ Step 8: First Draft — {chapter_name}")
    run_scene_draft(config, project_dir, chapter_name, slug, previous_path, skip_edit)
    return slug


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="AI Novel Writing Automation Tool")
    parser.add_argument("--project", help="Override project directory from config")
    parser.add_argument("--provider", choices=["openai", "anthropic", "openrouter"])
    parser.add_argument("--model")
    parser.add_argument("--skip-edit", action="store_true")

    subparsers = parser.add_subparsers(dest="command", required=True)
    subparsers.add_parser("init", help="Creates project folder structure and blank config")
    subparsers.add_parser("run", help="Runs all steps in sequence")

    step_parser = subparsers.add_parser("step", help="Run only step N")
    step_parser.add_argument("number", type=int)

    scene_parser = subparsers.add_parser("scene", help="Run scene brief + first draft for chapter")
    scene_parser.add_argument("chapter_name")

    subparsers.add_parser("list", help="List all steps and completion status")
    return parser.parse_args()


def initialize_project(project_dir: Path) -> None:
    (project_dir / "inputs").mkdir(parents=True, exist_ok=True)
    (project_dir / "outputs/scenes").mkdir(parents=True, exist_ok=True)

    config_path = project_dir / "config.json"
    if not config_path.exists():
        data = DEFAULT_CONFIG.copy()
        data["project_dir"] = str(project_dir)
        save_json(config_path, data)

    progress_path = project_dir / "progress.json"
    if not progress_path.exists():
        save_json(progress_path, DEFAULT_PROGRESS)

    placeholders = [
        "braindump.docx",
        "genre_tropes.docx",
        "outline_template.docx",
        "writing_samples.docx",
        "prohibited_words.docx",
    ]
    for name in placeholders:
        p = project_dir / "inputs" / name
        if not p.exists():
            Document().save(p)

    console.print(f"✅ Project initialized at: {project_dir}")


def build_runtime_config(args: argparse.Namespace) -> RuntimeConfig:
    project_dir = Path(args.project or DEFAULT_CONFIG["project_dir"]).resolve()
    config_path = project_dir / "config.json"
    config_data = load_json(config_path, DEFAULT_CONFIG)

    provider = args.provider or config_data.get("provider", "openai")
    model = args.model or config_data.get("model", "gpt-4o")

    runtime = RuntimeConfig(
        provider=provider,
        model=model,
        project_dir=Path(args.project or config_data.get("project_dir", project_dir)).resolve(),
        book_title=config_data.get("book_title", DEFAULT_CONFIG["book_title"]),
        genre=config_data.get("genre", DEFAULT_CONFIG["genre"]),
    )
    return runtime


def print_header(config: RuntimeConfig) -> None:
    console.print(
        Panel.fit(
            "🖊  AI Novel Writing Automation Tool\n"
            "====================================\n"
            f"Project: {config.project_dir}\n"
            f"Provider: {config.provider} ({config.model})\n"
            f"Book: {config.book_title} | Genre: {config.genre}"
        )
    )


def main() -> None:
    load_dotenv()
    args = parse_args()

    if args.command == "init":
        project_dir = Path(args.project or DEFAULT_CONFIG["project_dir"]).resolve()
        initialize_project(project_dir)
        return

    config = build_runtime_config(args)
    initialize_project(config.project_dir)
    print_header(config)

    progress_path = config.project_dir / "progress.json"
    progress = load_json(progress_path, DEFAULT_PROGRESS)

    try:
        if args.command == "list":
            for step_num in range(1, 7):
                key = f"step{step_num}_{STEP_SPECS[step_num]['name'].lower().replace(' ', '_')}"
                file_exists = (config.project_dir / STEP_SPECS[step_num]["output"]).exists()
                status = "✅ complete" if progress.get(key) or file_exists else "⬜ pending"
                console.print(f"Step {step_num}: {STEP_SPECS[step_num]['name']} — {status}")
            scenes = progress.get("scenes", {})
            if scenes:
                console.print("Scenes:")
                for slug, scene_status in scenes.items():
                    console.print(
                        f"- {slug}: brief={'✅' if scene_status.get('brief') else '⬜'}, "
                        f"draft={'✅' if scene_status.get('draft') else '⬜'}"
                    )
            else:
                console.print("Scenes: none tracked")
            return

        if args.command == "step":
            n = args.number
            if n < 1 or n > 8:
                console.print("[red]Step must be between 1 and 8[/red]")
                raise SystemExit(1)
            if n <= 6:
                execute_step(n, config, config.project_dir, args.skip_edit)
                progress[f"step{n}_{STEP_SPECS[n]['name'].lower().replace(' ', '_')}"] = True
            else:
                chapter_name = input("Enter chapter name (e.g. 'Chapter 1: The Awakening'): ").strip()
                slug = scene_flow(config, config.project_dir, chapter_name, args.skip_edit)
                progress.setdefault("scenes", {}).setdefault(slug, {})
                progress["scenes"][slug]["brief"] = True
                progress["scenes"][slug]["draft"] = True
            save_json(progress_path, progress)
            return

        if args.command == "scene":
            slug = scene_flow(config, config.project_dir, args.chapter_name, args.skip_edit)
            progress.setdefault("scenes", {}).setdefault(slug, {})
            progress["scenes"][slug]["brief"] = True
            progress["scenes"][slug]["draft"] = True
            save_json(progress_path, progress)
            return

        if args.command == "run":
            for step_num in range(1, 7):
                out_path = config.project_dir / STEP_SPECS[step_num]["output"]
                if out_path.exists():
                    should_skip = input(
                        f"Output exists for Step {step_num} ({out_path.relative_to(config.project_dir)}). Skip? (y/n): "
                    ).strip().lower()
                    if should_skip == "y":
                        progress[f"step{step_num}_{STEP_SPECS[step_num]['name'].lower().replace(' ', '_')}"] = True
                        continue

                execute_step(step_num, config, config.project_dir, args.skip_edit)
                progress[f"step{step_num}_{STEP_SPECS[step_num]['name'].lower().replace(' ', '_')}"] = True
                save_json(progress_path, progress)

            console.print("\n▶ Scenes")
            while True:
                chapter_name = input("Enter chapter name (e.g. 'Chapter 1: The Awakening'): ").strip()
                if not chapter_name:
                    break
                slug = scene_flow(config, config.project_dir, chapter_name, args.skip_edit)
                progress.setdefault("scenes", {}).setdefault(slug, {})
                progress["scenes"][slug]["brief"] = True
                progress["scenes"][slug]["draft"] = True
                save_json(progress_path, progress)

                more = input("Add another chapter? (y/n): ").strip().lower()
                if more != "y":
                    break

    except KeyboardInterrupt:
        save_json(progress_path, progress)
        console.print("\n[yellow]Stopping. Progress saved.[/yellow]")


if __name__ == "__main__":
    main()
